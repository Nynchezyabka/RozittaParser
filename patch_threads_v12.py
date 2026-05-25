"""
Патч v12 — дедупликация + дерево ответов в thread-режиме.

Исправляет v11: _replace_method теперь ищет метод по границам класса
(находит class Xxx, затем следующий class — и ищет def внутри).
Это работает даже если между class и def сотни строк.

Проблемы, решаемые патчем:
  1. Сообщения дублируются — одна пара рендерит msg как context, следующая как reply
  2. HTML без отступов — всё в столбик, не выглядит ветками форума

Решение:
  - _dedup_thread_messages(pairs) — собирает уникальные сообщения,
    вычисляет depth (глубину в дереве ответов), убирает дубликаты
  - Docx: линейный рендер с маркером "↳ в ответ на [автор]"
  - Markdown: линейный рендер с маркером "↳"
  - HTML: дерево с CSS-отступами (depth-0..depth-5)
  - JSON: плоский список с полями depth + reply_to_author

Идемпотентность: патч проверяет, были ли изменения уже применены (v11).

Запуск:
    cd D:\\RozittaParser\\RozittaParser
    python patch_threads_v12.py
"""
import os
import re

BASE = r"D:\RozittaParser\RozittaParser"
PATCHES_APPLIED = 0
PATCHES_TOTAL = 0


def _patch(rel_path: str, old: str, new: str, label: str) -> None:
    global PATCHES_APPLIED, PATCHES_TOTAL
    PATCHES_TOTAL += 1
    fpath = os.path.join(BASE, rel_path.replace("/", os.sep))
    if not os.path.isfile(fpath):
        print(f"  [SKIP] {label}: файл не найден")
        return
    text = open(fpath, encoding="utf-8").read()
    old_n = old.replace("\r\n", "\n")
    new_n = new.replace("\r\n", "\n")
    text_n = text.replace("\r\n", "\n")
    if new_n in text_n:
        # Уже применено
        print(f"  [SKIP] {label}: уже применено")
        PATCHES_APPLIED += 1
        return
    if old_n not in text_n:
        print(f"  [MISS] {label}")
        return
    text = text_n.replace(old_n, new_n, 1)
    open(fpath, "w", encoding="utf-8").write(text)
    PATCHES_APPLIED += 1
    print(f"  [ OK ] {label}")


def _insert_after(fpath: str, after_text: str, insert_text: str, label: str) -> None:
    """Вставляет insert_text после after_text в файле. Идемпотентна."""
    global PATCHES_APPLIED, PATCHES_TOTAL
    PATCHES_TOTAL += 1
    if not os.path.isfile(fpath):
        print(f"  [SKIP] {label}: файл не найден")
        return
    text = open(fpath, encoding="utf-8").read().replace("\r\n", "\n")
    # Сначала проверяем — не была ли уже вставка сделана
    # Используем уникальный маркер из insert_text (имя функции)
    _marker = None
    for line in insert_text.strip().splitlines():
        if line.strip().startswith("def "):
            _marker = line.strip()
            break
    if _marker and _marker in text:
        print(f"  [SKIP] {label}: уже применено")
        PATCHES_APPLIED += 1
        return
    after_n = after_text.replace("\r\n", "\n")
    if after_n not in text:
        print(f"  [MISS] {label}: якорь не найден")
        return
    text = text.replace(after_n, after_n + insert_text, 1)
    open(fpath, "w", encoding="utf-8").write(text)
    PATCHES_APPLIED += 1
    print(f"  [ OK ] {label}")


def _replace_method(fpath: str, class_name: str, method_name: str,
                    new_body: str, label: str) -> None:
    """
    Программно заменяет тело метода в классе.

    Алгоритм v12:
    1. Находим строку `class ClassName`
    2. Находим следующий `class ` после неё (граница класса)
    3. Ищем `def method_name` между class_start и class_end
    4. Находим конец метода (следующий def на том же/меньшем уровне отступа)
    5. Заменяем метод на new_body (или удаляем, если new_body пустой)
    """
    global PATCHES_APPLIED, PATCHES_TOTAL
    PATCHES_TOTAL += 1
    if not os.path.isfile(fpath):
        print(f"  [SKIP] {label}: файл не найден")
        return
    text = open(fpath, encoding="utf-8").read().replace("\r\n", "\n")
    lines = text.split("\n")

    # ── Шаг 1: Найти начало класса ─────────────────────────────────
    class_start = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith(f"class {class_name}") and stripped.endswith(":"):
            class_start = i
            break
    if class_start is None:
        print(f"  [MISS] {label}: класс {class_name} не найден")
        return

    # ── Шаг 2: Найти конец класса (следующий class на том же уровне) ──
    class_end = len(lines)
    for i in range(class_start + 1, len(lines)):
        stripped = lines[i].strip()
        if stripped.startswith("class ") and stripped.endswith(":"):
            # Проверяем что это top-level класс (без отступа)
            indent = len(lines[i]) - len(lines[i].lstrip())
            if indent == 0:
                class_end = i
                break

    # ── Шаг 3: Найти метод внутри класса ────────────────────────────
    method_start = None
    for i in range(class_start + 1, class_end):
        if f"def {method_name}" in lines[i]:
            method_start = i
            break
    if method_start is None:
        print(f"  [MISS] {label}: метод {class_name}.{method_name} не найден")
        return

    # ── Шаг 4: Найти конец метода ───────────────────────────────────
    base_indent = len(lines[method_start]) - len(lines[method_start].lstrip())
    method_end = class_end  # по умолчанию — до конца класса
    for i in range(method_start + 1, class_end):
        stripped = lines[i].strip()
        if not stripped:
            continue
        line_indent = len(lines[i]) - len(lines[i].lstrip())
        if line_indent <= base_indent and stripped.startswith("def "):
            method_end = i
            break

    # ── Шаг 5: Заменить / удалить метод ─────────────────────────────
    before = lines[:method_start]
    after = lines[method_end:]

    if new_body.strip():
        # Замена тела метода
        new_lines = before + new_body.rstrip("\n").split("\n") + after
    else:
        # Удаление метода — убираем пустые строки перед ним
        while before and before[-1].strip() == "":
            before.pop()
        new_lines = before + after

    result = "\n".join(new_lines)
    open(fpath, "w", encoding="utf-8").write(result)
    PATCHES_APPLIED += 1
    print(f"  [ OK ] {label} (строки {method_start+1}\u2013{method_end})")


# ============================================================
print("=" * 60)
print("Патч v12 — дедупликация + дерево ответов")
print("=" * 60)

GEN_PATH = os.path.join(BASE, "features", "export", "generator.py")

# ── 1. Добавить _dedup_thread_messages после _topic_suffix ──
print("\n\U0001f4e6 1. Добавить _dedup_thread_messages()")

_dedup_code = r'''

def _dedup_thread_messages(pairs: list) -> list:
    """
    Преобразует пары (context→reply) в дедуплицированный список сообщений
    с информацией о глубине в дереве ответов.

    Устраняет дубликаты: каждое сообщение появляется ровно один раз.
    Вычисляет depth (0 = корень) по цепочке reply_to_msg_id.

    Returns:
        Список кортежей (row, depth, reply_to_author), отсортированных по дате.
        row             — tuple из БД
        depth           — int: 0 для корневого сообщения, +1 за каждый уровень
        reply_to_author — str|None: автор сообщения-родителя
    """
    seen_ids: set = set()
    messages: list = []
    msg_by_id: dict = {}

    for pair in pairs:
        for row in pair["context"]:
            mid = row[_COL_MESSAGE_ID]
            if mid not in seen_ids:
                seen_ids.add(mid)
                messages.append(row)
                msg_by_id[mid] = row
        for row in pair["reply"]:
            mid = row[_COL_MESSAGE_ID]
            if mid not in seen_ids:
                seen_ids.add(mid)
                messages.append(row)
                msg_by_id[mid] = row

    # Сортировка по дате
    messages.sort(key=lambda r: r[_COL_DATE])

    # Вычисление depth по цепочке reply_to
    depth_cache: dict = {}

    def _depth(row) -> int:
        mid = row[_COL_MESSAGE_ID]
        if mid in depth_cache:
            return depth_cache[mid]
        reply_to = row[_COL_REPLY_TO]
        if reply_to is None or reply_to not in msg_by_id:
            depth_cache[mid] = 0
            return 0
        parent_depth = _depth(msg_by_id[reply_to])
        depth_cache[mid] = parent_depth + 1
        return depth_cache[mid]

    result: list = []
    for row in messages:
        depth = _depth(row)
        reply_to = row[_COL_REPLY_TO]
        reply_author = None
        if reply_to and reply_to in msg_by_id:
            reply_author = (
                msg_by_id[reply_to][_COL_USERNAME]
                or f"id_{msg_by_id[reply_to][_COL_USER_ID]}"
            )
        result.append((row, depth, reply_author))

    return result
'''

_insert_after(
    GEN_PATH,
    after_text='    return f"_topic{topic_id}" if topic_id is not None else ""\n',
    insert_text=_dedup_code,
    label="_dedup_thread_messages() после _topic_suffix()",
)

# ── 2. DocxGenerator._generate_threads — линейный с ↳ маркерами ──
print("\n\U0001f4e6 2. DocxGenerator._generate_threads")

_docx_threads = '''\
    def _generate_threads(
        self,
        chat_id:      int,
        chat_title:   str,
        topic_id:     Optional[int],
        topic_name:   Optional[str],
        user_id:      int,
        username:     Optional[str],
        period_label: str,
        date_from:    Optional[str],
        date_to:      Optional[str],
        log:          "_LogCallback",
    ) -> List[str]:
        """Создаёт DOCX с деревом ответов для thread-режима (без дубликатов)."""
        from docx import Document
        from docx.shared import Pt
        from core.utils import sanitize_filename

        doc = Document()
        user_label   = sanitize_filename(username or f"id_{user_id}")
        name_parts   = [sanitize_filename(chat_title)]
        if topic_name:
            name_parts.append(sanitize_filename(topic_name))
        name_parts  += [user_label, "threads", period_label]
        out_path     = Path(self._output_dir) / ("_".join(name_parts) + ".docx")

        doc.add_heading(
            f"{chat_title} \u2014 ветки с {username or user_label}",
            level=1,
        )

        pairs = self._db.get_thread_pairs(
            chat_id, user_id,
            topic_id=topic_id, date_from=date_from, date_to=date_to,
        )
        if log:
            log(f"Thread DOCX: {len(pairs)} пар")

        deduped = _dedup_thread_messages(pairs)
        if log:
            log(f"Thread DOCX: {len(deduped)} уникальных сообщений")

        for row, depth, reply_author in deduped:
            indent_pt = Pt(24 * depth)
            author     = row[_COL_USERNAME] or f"id_{row[_COL_USER_ID]}"
            date_str   = (row[_COL_DATE] or "")[:16]
            text       = (row[_COL_TEXT] or "").strip()

            # Заголовок
            header_p = doc.add_paragraph()
            header_p.paragraph_format.left_indent = indent_pt
            header_p.paragraph_format.space_after  = Pt(0)

            prefix = "\u21b3 " if depth > 0 else ""
            run = header_p.add_run(f"{prefix}{author}")
            run.bold = True
            run.font.size = Pt(11) if depth == 0 else Pt(10)
            run.font.color.rgb = RGBColor(51, 51, 51) if depth == 0 else RGBColor(80, 80, 80)

            date_run = header_p.add_run(f"  {date_str}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(128, 128, 128)

            # Маркер ответа
            if depth > 0 and reply_author:
                reply_p = doc.add_paragraph()
                reply_p.paragraph_format.left_indent = indent_pt
                reply_p.paragraph_format.space_before = Pt(0)
                reply_p.paragraph_format.space_after  = Pt(2)
                rr = reply_p.add_run(f"\u21a9 в ответ на: {reply_author}")
                rr.font.italic = True
                rr.font.size = Pt(8)
                rr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # Текст
            if text:
                text_p = doc.add_paragraph()
                text_p.paragraph_format.left_indent = indent_pt
                text_p.paragraph_format.space_before = Pt(2)
                xml_magic.write_text_with_links(text_p, text)

            # Разделитель (только для depth=0)
            if depth == 0:
                sep = doc.add_paragraph("\u2500" * 42)
                sep.paragraph_format.space_before = Pt(3)
                sep.paragraph_format.space_after  = Pt(3)

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return [str(out_path)]
'''

_replace_method(GEN_PATH, "DocxGenerator", "_generate_threads", _docx_threads,
               label="DocxGenerator._generate_threads \u2192 dedup + \u21b3")

# ── 3. Удалить _add_context_block_to_doc (больше не нужен) ──
print("\n\U0001f4e6 3. Удалить DocxGenerator._add_context_block_to_doc")

_replace_method(GEN_PATH, "DocxGenerator", "_add_context_block_to_doc", "",
               label="DocxGenerator._add_context_block_to_doc \u2192 удалить")

# ── 4. MarkdownGenerator._generate_threads ──
print("\n\U0001f4e6 4. MarkdownGenerator._generate_threads")

_md_threads = '''\
    def _generate_threads(
        self,
        chat_id:      int,
        chat_title:   str,
        topic_id:     Optional[int],
        topic_name:   Optional[str],
        user_id:      int,
        username:     Optional[str],
        period_label: str,
        date_from:    Optional[str],
        date_to:      Optional[str],
        log:          "_LogCallback",
    ) -> List[str]:
        """Markdown-выгрузка для thread-режима (без дубликатов)."""
        from core.utils import sanitize_filename

        stt_map: dict[int, str] = {}
        try:
            stt_map = self._db.get_transcriptions_for_chat(chat_id)
        except Exception:
            pass

        pairs = self._db.get_thread_pairs(
            chat_id, user_id,
            topic_id=topic_id, date_from=date_from, date_to=date_to,
        )
        if log:
            log(f"Thread MD: {len(pairs)} пар")

        deduped = _dedup_thread_messages(pairs)
        if log:
            log(f"Thread MD: {len(deduped)} уникальных сообщений")

        user_label = sanitize_filename(username or f"id_{user_id}")
        name_parts = [sanitize_filename(chat_title)]
        if topic_name:
            name_parts.append(sanitize_filename(topic_name))
        name_parts += [user_label, "threads", period_label]
        out_path = Path(self._output_dir) / ("_".join(name_parts) + ".md")
        out_path.parent.mkdir(parents=True, exist_ok=True)

        user_display = username or user_label
        lines = [f"# {chat_title} \u2014 ветки с {user_display}\n\n"]

        for row, depth, reply_author in deduped:
            raw_date = row[_COL_DATE] or ""
            date_str = raw_date[:16].replace("T", " ") if raw_date else "\u2014"
            author   = row[_COL_USERNAME] or f"id:{row[_COL_USER_ID]}" or "Неизвестно"
            text     = (row[_COL_TEXT] or "").strip()
            indent   = "    " * depth

            # Заголовок
            prefix = "\u21b3 " if depth > 0 else ""
            line = f"{indent}**[{date_str}] {prefix}{author}:**"

            # Маркер ответа
            if depth > 0 and reply_author:
                line += f"  *(в ответ на: {reply_author})*"

            lines.append(line)
            if text:
                lines.append(f"{indent}{text}")

            # STT
            msg_id = row[_COL_MESSAGE_ID]
            stt = stt_map.get(msg_id)
            if stt:
                lines.append(f"{indent}*(STT: {stt.strip()})*")

            if depth == 0:
                lines.append("\n---\n")

        out_path.write_text("\n".join(lines), encoding="utf-8")
        return [str(out_path)]
'''

_replace_method(GEN_PATH, "MarkdownGenerator", "_generate_threads", _md_threads,
               label="MarkdownGenerator._generate_threads \u2192 dedup")

# ── 5. Удалить MarkdownGenerator._format_thread_pair ──
print("\n\U0001f4e6 5. Удалить MarkdownGenerator._format_thread_pair")

_replace_method(GEN_PATH, "MarkdownGenerator", "_format_thread_pair", "",
               label="MarkdownGenerator._format_thread_pair \u2192 удалить")

# ── 6. HtmlGenerator._generate_threads — дерево с отступами ──
print("\n\U0001f4e6 6. HtmlGenerator._generate_threads")

_html_threads = '''\
    def _generate_threads(
        self,
        chat_id:      int,
        chat_title:   str,
        topic_id:     Optional[int],
        topic_name:   Optional[str],
        user_id:      int,
        username:     Optional[str],
        period_label: str,
        date_from:    Optional[str],
        date_to:      Optional[str],
        log:          "_LogCallback",
    ) -> List[str]:
        """HTML-выгрузка для thread-режима (дерево с отступами, без дубликатов)."""
        from core.utils import sanitize_filename

        stt_map: dict[int, str] = {}
        try:
            stt_map = self._db.get_transcriptions_for_chat(chat_id)
        except Exception:
            pass

        pairs = self._db.get_thread_pairs(
            chat_id, user_id,
            topic_id=topic_id, date_from=date_from, date_to=date_to,
        )
        if log:
            log(f"Thread HTML: {len(pairs)} пар")

        deduped = _dedup_thread_messages(pairs)
        if log:
            log(f"Thread HTML: {len(deduped)} уникальных сообщений")

        # Рендерим сообщения напрямую через _HTML_MSG
        body_parts: list = []
        for row, depth, reply_author in deduped:
            msg_id     = row[_COL_MESSAGE_ID]
            author     = row[_COL_USERNAME] or f"id_{row[_COL_USER_ID]}" or "Неизвестно"
            raw_date   = row[_COL_DATE] or ""
            date_str   = raw_date[:16].replace("T", " ") if raw_date else "\u2014"
            text       = (row[_COL_TEXT] or "").strip()
            media_path = row[_COL_MEDIA_PATH]
            file_type  = row[_COL_FILE_TYPE] or ""
            avatar_letter = (author[0] or "?").upper()

            safe_text = html_lib.escape(text)
            safe_text = re.sub(
                r"(https?://\\S+)",
                r'<a href="\\1">\\1</a>',
                safe_text,
            )
            text_block = f'<div class="msg-text">{safe_text}</div>' if text else ""

            media_block = ""
            if media_path and os.path.exists(media_path):
                abs_path = os.path.abspath(media_path)
                fname = os.path.basename(abs_path)
                if is_image_path(abs_path):
                    media_block = (
                        f'<img class="msg-img" src="{html_lib.escape(abs_path)}" '
                        f'alt="{html_lib.escape(fname)}">'
                    )
                else:
                    media_block = (
                        f'<div class="msg-media"><a href="{html_lib.escape(abs_path)}">'
                        f'{html_lib.escape(fname)}</a></div>'
                    )

            reply_badge = ""
            if depth > 0 and reply_author:
                reply_badge = (
                    f'<span class="reply-badge">\u21b3 в ответ на: '
                    f'{html_lib.escape(reply_author)}</span>'
                )

            stt_block = ""
            stt = stt_map.get(msg_id)
            if stt:
                stt_block = (
                    f'<div class="msg-stt">\U0001f399 {html_lib.escape(stt.strip())}</div>'
                )

            depth_class = f"depth-{min(depth, 5)}"

            body_parts.append(_HTML_MSG.format(
                depth=depth,
                msg_id=msg_id,
                avatar_letter=avatar_letter,
                author=html_lib.escape(author),
                date=date_str,
                reply_badge=reply_badge,
                quote_block="",
                text_block=text_block,
                media_block=media_block,
                stt_block=stt_block,
            ))

        body = "\n".join(body_parts)
        user_label = sanitize_filename(username or f"id_{user_id}")
        name_parts = [sanitize_filename(chat_title)]
        if topic_name:
            name_parts.append(sanitize_filename(topic_name))
        name_parts += [user_label, "threads", period_label]
        out_path = Path(self._output_dir) / ("_".join(name_parts) + ".html")
        out_path.parent.mkdir(parents=True, exist_ok=True)

        h_title = html_lib.escape(
            f"{chat_title} \u2014 ветки с {username or user_label}"
        )
        html = _HTML_TEMPLATE.format(
            title=h_title,
            total=len(deduped),
            body=body,
        )
        out_path.write_text(html, encoding="utf-8")
        return [str(out_path)]
'''

_replace_method(GEN_PATH, "HtmlGenerator", "_generate_threads", _html_threads,
               label="HtmlGenerator._generate_threads \u2192 дерево с отступами")

# ── 7. JsonGenerator._generate_threads ──
print("\n\U0001f4e6 7. JsonGenerator._generate_threads")

_json_threads = '''\
    def _generate_threads(
        self,
        chat_id:      int,
        chat_title:   str,
        topic_id:     Optional[int],
        topic_name:   Optional[str],
        user_id:      int,
        username:     Optional[str],
        period_label: str,
        date_from:    Optional[str],
        date_to:      Optional[str],
        log:          "_LogCallback",
    ) -> List[str]:
        """
        JSON-выгрузка для thread-режима (без дубликатов, с depth).

        Каждая запись содержит:
          - все поля обычного сообщения
          - depth: глубина в дереве ответов (0 = корень)
          - reply_to_author: автор сообщения-родителя (если есть)
        """
        import json
        from core.utils import sanitize_filename

        stt_map: dict[int, str] = {}
        try:
            stt_map = self._db.get_transcriptions_for_chat(chat_id)
        except Exception:
            pass

        pairs = self._db.get_thread_pairs(
            chat_id, user_id,
            topic_id=topic_id, date_from=date_from, date_to=date_to,
        )
        if log:
            log(f"Thread JSON: {len(pairs)} пар")

        deduped = _dedup_thread_messages(pairs)
        if log:
            log(f"Thread JSON: {len(deduped)} уникальных сообщений")

        records = []
        for row, depth, reply_author in deduped:
            msg_id = row[_COL_MESSAGE_ID]
            rec = self._make_record(row, stt_map.get(msg_id))
            rec["depth"] = depth
            rec["reply_to_author"] = reply_author
            rec["type"] = "thread_reply" if depth > 0 else "thread_root"
            records.append(rec)

        user_label = sanitize_filename(username or f"id_{user_id}")
        name_parts = [sanitize_filename(chat_title)]
        if topic_name:
            name_parts.append(sanitize_filename(topic_name))
        name_parts += [user_label, "threads", period_label]
        out_path = Path(self._output_dir) / ("_".join(name_parts) + ".json")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(
            json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        return [str(out_path)]
'''

_replace_method(GEN_PATH, "JsonGenerator", "_generate_threads", _json_threads,
               label="JsonGenerator._generate_threads \u2192 dedup + depth")

# ── 8. Добавить CSS для depth-4, depth-5 в HTML шаблон ──
print("\n\U0001f4e6 8. HTML: добавить CSS для depth-4, depth-5")

# Проверяем — возможно уже добавлено v11
_patch(
    "features/export/generator.py",
    old='  .depth-3 {{ margin-left: 80px; }}\n',
    new='  .depth-3 {{ margin-left: 80px; }}\n  .depth-4 {{ margin-left: 104px; }}\n  .depth-5 {{ margin-left: 128px; }}\n',
    label="CSS: +depth-4, depth-5",
)

# ── 9. DocxGenerator.generate(): убрать дублирующийся вызов _generate_threads ──
print("\n\U0001f4e6 9. DocxGenerator.generate(): убрать дублирующий threads dispatch")

# Проверяем — возможно уже удалено v11
_patch(
    "features/export/generator.py",
    old="""\
        try:
            if user_filter_mode == "threads":
                files = self._generate_threads(
                    chat_id          = chat_id,
                    chat_title       = chat_title or f"chat_{chat_id}",
                    topic_id         = topic_id,
                    topic_name       = topic_name,
                    user_id          = user_id,
                    username         = username,
                    period_label     = period_label,
                    date_from        = date_from,
                    date_to          = date_to,
                    log              = self._log,
                )
            elif split_mode == "post":""",
    new="""\
        try:
            if split_mode == "post":""",
    label="DocxGenerator.generate(): убрать дублирующий threads dispatch",
)

# ============================================================
print(f"\n{'=' * 60}")
print(f"Применено изменений: {PATCHES_APPLIED}/{PATCHES_TOTAL}")
print("=" * 60)

if PATCHES_APPLIED == PATCHES_TOTAL:
    print("\n\u2705 Все патчи применены!")
else:
    print(f"\n\u26a0\ufe0f  Применено {PATCHES_APPLIED} из {PATCHES_TOTAL}.")

print("""
  Что изменилось:
  ─────────────
  Новая функция _dedup_thread_messages(pairs):
    \u2022 Собирает все уникальные сообщения из всех пар
    \u2022 Вычисляет depth (глубину в дереве ответов)
    \u2022 Определяет reply_to_author для маркеров
    \u2022 Возвращает [(row, depth, reply_author), ...] без дубликатов

  DocxGenerator._generate_threads:
    \u2022 Линейный рендер, depth \u2192 отступ (24pt \u00d7 depth)
    \u2022 depth>0: маркер "\u21b3" + "\u21a9 в ответ на: [автор]"
    \u2022 Разделитель только между корневыми сообщениями (depth=0)

  MarkdownGenerator._generate_threads:
    \u2022 Линейный рендер, depth \u2192 отступ (4 пробела \u00d7 depth)
    \u2022 depth>0: маркер "\u21b3" + "*(в ответ на: автор)*"

  HtmlGenerator._generate_threads:
    \u2022 Дерево с CSS-классами depth-0..depth-5
    \u2022 Каждый reply показывает "\u21b3 в ответ на: [автор]"
    \u2022 CSS margin-left нарастает с глубиной

  JsonGenerator._generate_threads:
    \u2022 Плоский список с полями depth + reply_to_author
    \u2022 type: "thread_root" | "thread_reply"

  Удалены устаревшие методы:
    \u2022 DocxGenerator._add_context_block_to_doc
    \u2022 MarkdownGenerator._format_thread_pair

  Исправление v11 \u2192 v12:
    \u2022 _replace_method теперь ищет метод по границам класса
      (находит class Xxx, затем следующий class, и ищет def внутри)
    \u2022 Это работает даже если между class и def сотни строк
""")
