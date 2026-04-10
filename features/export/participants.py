"""
features/export/participants.py — Экспорт списка участников.

Поддерживает два формата:
  - DOCX (рекомендуемый): таблица с активными ссылками на профили Telegram.
  - MD (резервный): простой текстовый список для просмотра/копирования.

Нет импортов Qt. Нет Telethon. Только stdlib + python-docx + core.utils.

Публичный API:
    get_user_message_counts(db_path, chat_id) → dict[int, int]
        Читает счётчики из локальной SQLite (резерв если API недоступен).

    enrich_and_sort_users(users, counts) → list[dict]
        Добавляет msg_count и сортирует по убыванию активности.

    export_participants_docx(users, chat_title, output_dir, counts) → str
        Создаёт DOCX с таблицей и кликабельными ссылками на профили.

    export_participants_md(users, chat_title, output_dir, counts) → str
        Создаёт MD-файл (резервный формат).
"""

from __future__ import annotations

import logging
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional

from core.utils import sanitize_filename

logger = logging.getLogger(__name__)


# ==============================================================================
# Счётчики из локальной БД (резервный источник)
# ==============================================================================

def get_user_message_counts(db_path: str, chat_id: int) -> Dict[int, int]:
    """
    Возвращает {user_id: msg_count} из локальной SQLite.

    Используется как fallback если get_user_stats недоступен.
    Основной источник — поле message_count в словарях от MembersWorker.
    """
    if not db_path or not os.path.exists(db_path):
        return {}
    try:
        conn = sqlite3.connect(db_path, timeout=5, check_same_thread=False)
        conn.execute("PRAGMA journal_mode=WAL")
        cursor = conn.execute(
            """
            SELECT user_id, COUNT(*) AS msg_count
            FROM messages
            WHERE chat_id = ?
              AND user_id IS NOT NULL
              AND user_id != 0
              AND is_comment = 0
            GROUP BY user_id
            ORDER BY msg_count DESC
            """,
            (chat_id,),
        )
        result = {int(r[0]): int(r[1]) for r in cursor.fetchall()}
        conn.close()
        return result
    except Exception as exc:
        logger.warning("participants: get_user_message_counts failed: %s", exc)
        return {}


# ==============================================================================
# Утилиты
# ==============================================================================

def _display_name(user: dict) -> str:
    uid = user.get("id", 0)
    return (
        user.get("name")
        or user.get("username")
        or user.get("first_name")
        or str(uid)
    )


def _msg_count(user: dict, counts: Optional[Dict[int, int]]) -> int:
    """Возвращает счётчик сообщений: сначала из поля словаря, потом из counts."""
    # get_user_stats уже возвращает message_count в словаре
    mc = user.get("message_count")
    if mc is not None:
        return int(mc)
    if counts:
        return counts.get(user.get("id", 0), 0)
    return 0


def enrich_and_sort_users(
    users:  List[dict],
    counts: Optional[Dict[int, int]] = None,
) -> List[dict]:
    """
    Добавляет поле msg_count и сортирует: по убыванию активности, затем по имени.

    Источник счётчика (в порядке приоритета):
      1. user["message_count"]  — поле от ChatsService.get_user_stats()
      2. counts[user_id]        — из локальной БД (get_user_message_counts)
      3. 0                      — нет данных
    """
    enriched = [{**u, "msg_count": _msg_count(u, counts)} for u in users]
    enriched.sort(key=lambda u: (-u["msg_count"], _display_name(u).lower()))
    return enriched


# ==============================================================================
# DOCX экспорт
# ==============================================================================

def export_participants_docx(
    users:      List[dict],
    chat_title: str,
    output_dir: str,
    counts:     Optional[Dict[int, int]] = None,
) -> str:
    """
    Создаёт DOCX-файл со списком участников.

    Таблица содержит:
      - #  (порядковый номер)
      - Имя (отображаемое, жирное)
      - @username (кликабельная ссылка tg://user?id=... или https://t.me/username)
      - Сообщений (если есть данные)

    Строки отсортированы по убыванию активности.

    Returns:
        Абсолютный путь к созданному .docx файлу.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from features.export import xml_magic

    os.makedirs(output_dir, exist_ok=True)

    now        = datetime.now()
    date_str   = now.strftime("%Y-%m-%d %H:%M")
    file_date  = now.strftime("%Y-%m-%d_%H-%M")
    safe_title = sanitize_filename(chat_title)
    filename   = f"{safe_title}_participants_{file_date}.docx"
    filepath   = os.path.join(output_dir, filename)

    has_counts = any(_msg_count(u, counts) > 0 for u in users)
    sorted_users = enrich_and_sort_users(users, counts)
    total_msgs   = sum(u["msg_count"] for u in sorted_users)

    doc = Document()

    # ── Стиль документа ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(11)

    # ── Заголовок ─────────────────────────────────────────────────────────
    title_p = doc.add_heading(f"Участники: {chat_title}", level=1)
    title_p.runs[0].font.color.rgb = RGBColor(0xFF, 0x6B, 0xC9)  # ACCENT_PINK

    # ── Метаданные ────────────────────────────────────────────────────────
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Дата выгрузки: ").bold = False
    meta_p.add_run(date_str).bold = True
    meta_p.add_run(f"   •   Участников: ").bold = False
    meta_p.add_run(str(len(sorted_users))).bold = True
    if has_counts and total_msgs:
        meta_p.add_run("   •   Сообщений в архиве: ").bold = False
        meta_p.add_run(f"{total_msgs:,}").bold = True
    meta_p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()  # пустая строка

    # ── Таблица ───────────────────────────────────────────────────────────
    col_count = 4 if has_counts else 3
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # Заголовки
    hdr_cells = table.rows[0].cells
    headers = ["#", "Имя", "Ссылка / @username"]
    if has_counts:
        headers.append("Сообщений")

    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x95, 0x00)  # ACCENT_ORANGE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, col_count - 1) else WD_ALIGN_PARAGRAPH.LEFT

    # Ширина столбцов
    col_widths = [Cm(1.0), Cm(5.5), Cm(7.5), Cm(2.5)][:col_count]
    for i, cell in enumerate(hdr_cells):
        cell.width = col_widths[i]

    # Строки участников
    for idx, user in enumerate(sorted_users, start=1):
        uid      = user.get("id", 0)
        name     = _display_name(user)
        username = user.get("username") or ""
        count    = user.get("msg_count", 0)

        row_cells = table.add_row().cells

        # #
        p_num = row_cells[0].paragraphs[0]
        p_num.add_run(str(idx))
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Имя
        p_name = row_cells[1].paragraphs[0]
        r = p_name.add_run(name)
        r.bold = True

        # Ссылка
        p_link = row_cells[2].paragraphs[0]
        if username:
            # Ссылка по username: открывается в браузере и в Telegram
            url        = f"https://t.me/{username}"
            link_text  = f"@{username}"
        else:
            # Ссылка по ID: открывает чат в Telegram Desktop/Mobile
            url       = f"tg://user?id={uid}"
            link_text = f"tg://user?id={uid}"

        xml_magic.add_external_hyperlink(p_link, url, link_text)

        # Сообщений
        if has_counts:
            p_cnt = row_cells[3].paragraphs[0]
            p_cnt.add_run(f"{count:,}" if count else "—")
            p_cnt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ширина столбцов — устанавливаем для каждой строки
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    doc.save(filepath)
    logger.info("participants: docx exported %d users → %s", len(sorted_users), filepath)
    return filepath


# ==============================================================================
# MD экспорт (резервный)
# ==============================================================================

def export_participants_md(
    users:      List[dict],
    chat_title: str,
    output_dir: str,
    counts:     Optional[Dict[int, int]] = None,
) -> str:
    """
    Создаёт Markdown-файл со списком участников (резервный формат).
    Для полноценного просмотра используйте export_participants_docx().
    """
    os.makedirs(output_dir, exist_ok=True)

    now        = datetime.now()
    date_str   = now.strftime("%Y-%m-%d %H:%M")
    file_date  = now.strftime("%Y-%m-%d_%H-%M")
    safe_title = sanitize_filename(chat_title)
    filename   = f"{safe_title}_participants_{file_date}.md"
    filepath   = os.path.join(output_dir, filename)

    sorted_users = enrich_and_sort_users(users, counts)
    has_counts   = any(u["msg_count"] > 0 for u in sorted_users)
    total_msgs   = sum(u["msg_count"] for u in sorted_users)

    stats_line = f"Всего участников: **{len(sorted_users)}**"
    if has_counts and total_msgs:
        stats_line += f"  |  Сообщений в архиве: **{total_msgs:,}**"

    lines: list[str] = [
        f"# Участники: {chat_title}", "",
        f"Дата выгрузки: {date_str}  ", stats_line, "",
        "---", "", "## Список участников", "",
    ]

    if has_counts:
        lines += ["| # | Имя | @username | Ссылка | Сообщений |",
                  "|--:|-----|-----------|--------|----------:|"]
    else:
        lines += ["| # | Имя | @username | Ссылка |",
                  "|--:|-----|-----------|--------|"]

    for idx, user in enumerate(sorted_users, start=1):
        uid      = user.get("id", 0)
        name     = _display_name(user).replace("|", "\\|")
        username = user.get("username") or ""
        count    = user.get("msg_count", 0)

        uname_md = f"@{username}" if username else "—"
        if username:
            link_md = f"[открыть](https://t.me/{username})"
        else:
            link_md = f"[открыть](tg://user?id={uid})"

        if has_counts:
            lines.append(f"| {idx} | {name} | {uname_md} | {link_md} | {count:,} |")
        else:
            lines.append(f"| {idx} | {name} | {uname_md} | {link_md} |")

    with open(filepath, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines) + "\n")

    logger.info("participants: md exported %d users → %s", len(sorted_users), filepath)
    return filepath